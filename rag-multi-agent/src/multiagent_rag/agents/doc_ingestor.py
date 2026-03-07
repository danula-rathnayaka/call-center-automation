import hashlib
import os

from docx import Document

from multiagent_rag.agents.base_ingestor import BaseIngestor
from multiagent_rag.utils.chunker import Chunker
from multiagent_rag.utils.logger import get_logger

logger = get_logger(__name__)


class DocIngestor(BaseIngestor):
    def __init__(self):
        self.chunker = Chunker()

    def process(self, file_path: str) -> tuple:
        try:
            doc = Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            combined_text = "\n".join(full_text)

            if not combined_text.strip():
                return [], ""

            doc_hash = self._compute_hash(file_path)

            metadata = {
                "source": os.path.basename(file_path),
                "type": "docx",
                "document_hash": doc_hash,
            }

            chunks = self.chunker.split_text(combined_text, metadata)
            return chunks, doc_hash

        except Exception as e:
            logger.error(f"Failed to process DOCX file {file_path}: {str(e)}")
            return [], ""

    def _compute_hash(self, file_path: str) -> str:
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for block in iter(lambda: f.read(8192), b""):
                sha256.update(block)
        return sha256.hexdigest()
